#!/usr/bin/env python3

"""
Professional Resume Generator
Creates a two-column resume document with formatting, icons, and hyperlinks
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
from docx2pdf import convert
import os, subprocess

def custom_print(msg):
    print(f"[CP] {msg}\n")

color_codes = {
    '1': 0x22,  # Dark Grey
    '2': 0x33,  # Darker Grey
    '3': 0x44,  # Grey
    '4': 0x55,  # Medium Grey
    '5': 0x88,  # Light Grey
    '6': 0xBB,  # Lighter Grey
    '7': 0xEE   # Very Light Grey
}
def grey(hex_component):
    """Returns an RGBColor with all components set to the same value."""
    return RGBColor(hex_component, hex_component, hex_component)

class ResumeGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_margins()
        
    def setup_margins(self):
        """Set document margins"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
    
    def resize_icon(self, icon_path, output_path, width_inch=0.15, height_inch=0.15):
        """Resize icon to specified dimensions"""
        try:
            img = Image.open(icon_path)
            # Convert inches to pixels (assuming 96 DPI)
            width_px = int(width_inch * 96)
            height_px = int(height_inch * 96)
            img_resized = img.resize((width_px, height_px), Image.Resampling.LANCZOS)
            img_resized.save(output_path)
            return output_path
        except Exception as e:
            print(f"Error resizing {icon_path}: {e}")
            return None
    
    def add_hyperlink(self, paragraph, url, text, color="#555555", underline=True):
        """
        Add a working clickable hyperlink to a paragraph (even inside tables).
        """
        # Get the document part (not the cell part)
        part = self.doc.part
        
        # Create the relationship
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        print(f"DEBUG: Created relationship {r_id} for URL: {url}")

        # Create hyperlink
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create new run
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Color
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color.lstrip('#'))
            rPr.append(c)

        # Underline
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        new_run.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        print(f"DEBUG: Hyperlink added to paragraph for text: {text}")


    def add_clickable_icon(self, paragraph, icon_path, url, width=0.21, height=0.21):
        """
        Add a clickable icon that links to a URL.
        """
        if not icon_path or not url:
            return
            
        # Resize icon
        resized_path = f"resized_{os.path.basename(icon_path)}"
        if not self.resize_icon(icon_path, resized_path, width, height):
            return
        
        # Use document part instead of paragraph part
        part = self.doc.part
        
        # Create relationship
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        print(f"DEBUG: Created icon relationship {r_id} for URL: {url}")
        
        # Add image first
        run = paragraph.add_run()
        run.add_picture(resized_path, width=Inches(width), height=Inches(height))
        
        # Get the drawing element from the run we just created
        drawing = run._element.find(qn('w:drawing'))
        
        if drawing is not None:
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create new run for the hyperlink
            new_run = OxmlElement('w:r')
            new_run.append(drawing)
            hyperlink.append(new_run)
            
            # Replace the old run with hyperlinked version
            paragraph._p.replace(run._element, hyperlink)
            print(f"DEBUG: Icon hyperlink added successfully")
        else:
            print(f"DEBUG: Could not find drawing element")
    
    def create_resume(self, data, icons):
        """
        Create the complete resume
        
        Args:
            data: Dictionary containing resume data
            icons: Dictionary containing icon file paths
        """
        # Create main table (2 columns)
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths (left column wider)
        table.columns[0].width = Inches(4.6)
        table.columns[1].width = Inches(3.0)
        
        # Add border to table
        self.set_table_borders(table) # renders black border lines (not good)
        
        # Get cells
        left_cell = table.rows[0].cells[0] # main content box
        right_cell = table.rows[0].cells[1] # address, location, etc box
        
        # Fill left column
        self.fill_left_column(left_cell, data)
        
        # Fill right column
        self.fill_right_column(right_cell, data, icons)
        
        return self.doc
    
    def set_table_borders(self, table):
        """Add borders to table"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def fill_left_column(self, cell, data):
        """Fill the left column with experience, projects, and education"""
        # Name and title at top
        p = cell.paragraphs[0]
        # p.style = p.part.document.styles['Normal']
        run = p.add_run(data['name'])
        run.bold = True
        # run.font.name = "Open Sans"
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Open Sans')  # ensure it’s recognized
        run.font.size = Pt(24) # Name size
        run.font.color.rgb = grey(color_codes['1'])
        p.paragraph_format.space_after = Pt(0) # No space after name
        
        p = cell.add_paragraph()
        run = p.add_run(data['title'])
        run.font.size = Pt(11)
        run.font.color.rgb = grey(color_codes['3'])
        p.paragraph_format.space_before = Pt(0) # Small space before title
        
        # # EXPERIENCE Section
        # self.add_section_header(cell, "EXPERIENCE")
        # EXPERIENCE Section
        p = cell.add_paragraph()               # Create the paragraph
        run = p.add_run("EXPERIENCE")          # Add text
        run.bold = True                        # Bold
        run.font.size = Pt(12)                 # Font size
        run.font.color.rgb = grey(color_codes['1'])  # Dark grey color
        p.paragraph_format.space_after = Pt(6) # Spacing after section header
        
        for exp in data['experience']:
            # Company name and location
            p = cell.add_paragraph()
            run = p.add_run(exp['company'])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_after = Pt(0)
            
            run = p.add_run(f" --- ")
            run = p.add_run(exp['role'])
            run.italic = True
            run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_after = Pt(0)
            
            # Date range
            p = cell.add_paragraph()
            run = p.add_run(exp['dates'])
            run.bold = True
            run.font.size = Pt(10)
            # run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
            # # Example for #6E6E6E
            # RGBColor(0x6E, 0x6E, 0x6E)
            run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_before = Pt(0)
            
            # # Description
            # p = cell.add_paragraph(exp['description'])
            # p.paragraph_format.space_after = Pt(12)
            # Description
            p = cell.add_paragraph()
            run = p.add_run(exp['description'])  # ✅ Create a run for the text
            # run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)  # ✅ Nice resume grey (#555555)
            run.font.size = Pt(11)  # Optional: fine-tune size for readability
            run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_after = Pt(12)
        
        # PROJECTS Section
        p = cell.add_paragraph()               # Create the paragraph
        run = p.add_run("PROJECTS")          # Add text
        run.bold = True                        # Bold
        run.font.size = Pt(12)                 # Font size
        run.font.color.rgb = grey(color_codes['1'])  # Dark grey color
        p.paragraph_format.space_after = Pt(6) # Spacing after section header
        
        # self.add_section_header(cell, "PROJECTS")
        
        for project in data['projects']:
            p = cell.add_paragraph()
            run = p.add_run(project['name'])
            run.bold = True
            run.font.color.rgb = grey(color_codes['1'])
            run.font.size = Pt(11)
            
            run = p.add_run(f" --- ")
            run = p.add_run(project['type'])
            run.font.color.rgb = grey(color_codes['1'])
            run.italic = True
            
            p = cell.add_paragraph()
            run = p.add_run(project['description'])  # ✅ Create a run for the text
            # run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)  # ✅ Nice resume grey (#555555)
            run.font.size = Pt(11)  # Optional: fine-tune size for readability
            run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_after = Pt(12)
        
            # p = cell.add_paragraph(project['description'])
            # p.paragraph_format.space_after = Pt(12)
        
        # # EDUCATION Section
        # self.add_section_header(cell, "EDUCATION")
        
        # for edu in data['education']:
        #     p = cell.add_paragraph()
        #     run = p.add_run(edu['institution'])
        #     run.bold = True
        #     run.font.size = Pt(11)
            
        #     run = p.add_run(f" --- ")
        #     run = p.add_run(edu['degree'])
        #     run.italic = True
            
        #     p = cell.add_paragraph()
        #     run = p.add_run(edu['year'])
        #     run.bold = True
        #     run.font.size = Pt(10)
    
    def fill_right_column(self, cell, data, icons):
        """Fill the right column with contact info, skills, awards, and languages"""
        # Contact Information with icons
        contact_info = [
            ('location', data['location'], None, 0.15, 0.15),
            ('phone', data['phone'], None, 0.15, 0.15),
            ('email', data['email'], f"mailto:{data['email']}", 0.17, 0.17),
            ('website', 'Portfolio Website', data['portfolio_url'], 0.15, 0.15),
        ]
        
        for icon_key, text, url, w, h in contact_info:
            p = cell.paragraphs[0] if icon_key == 'location' else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            
            # Add icon
            if icon_key in icons and icons[icon_key]:
                resized_path = f"resized_{icon_key}.png"
                self.resize_icon(icons[icon_key], resized_path, w, h)
                run = p.add_run()
                run.add_picture(resized_path, width=Inches(w), height=Inches(h))
            
            # Add text or hyperlink
            if url:
                self.add_hyperlink(p, url, text)
            else:
                run = p.add_run(f" {text}")
                run.font.color.rgb = grey(color_codes['1'])
                run.font.size = Pt(10)
        
        # Social media icons
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        
        social_media = [
            ('twitter', data.get('twitter_url')),
            ('linkedin', data.get('linkedin_url')),
            ('github', data.get('github_url')),
            ('whatsapp', data.get('whatsapp_url')),
        ]
        
        for icon_key, url in social_media:
            # if icon_key in icons and icons[icon_key]:
            #     resized_path = f"resized_{icon_key}.png"
            #     self.resize_icon(icons[icon_key], resized_path, 0.21, 0.21)
            #     run = p.add_run()
            #     run.add_picture(resized_path, width=Inches(0.21), height=Inches(0.21))
            #     run = p.add_run(" ")
            if icon_key in icons and icons[icon_key] and url:
                self.add_clickable_icon(p, icons[icon_key], url, 0.21, 0.21)
                # Add small space after each icon
                run = p.add_run(" ")
        
        # SKILLS Section
        p = cell.add_paragraph()               # Create the paragraph
        run = p.add_run("SKILLS")          # Add text
        run.bold = True                        # Bold
        run.font.size = Pt(12)                 # Font size
        run.font.color.rgb = grey(color_codes['1'])  # Dark grey color
        p.paragraph_format.space_after = Pt(6) # Spacing after section header
        
        # self.add_section_header(cell, "SKILLS")
        
        for skill in data['skills']:
            p = cell.add_paragraph(skill, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            # run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_after = Pt(6)
        
        # AWARDS Section
        p = cell.add_paragraph()               # Create the paragraph
        run = p.add_run("AWARDS")          # Add text
        run.bold = True                        # Bold
        run.font.size = Pt(12)                 # Font size
        run.font.color.rgb = grey(color_codes['1'])  # Dark grey color
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6) # Spacing after section header
        
        # self.add_section_header(cell, "AWARDS")
        
        for award in data['awards']:
            p = cell.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            # run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_after = Pt(6)
            
             # Add the title (bolded)
            title_run = p.add_run(award["title"])
            title_run.font.color.rgb = grey(color_codes['1'])
            title_run.bold = True  # Makes only the title bold

            # Add the rest of the text (not bold)
            year_run = p.add_run(f" -- {award['year']}")
            year_run.font.color.rgb = grey(color_codes['1'])
            year_run.bold = False # Keeps the year normal text
        
        # EDUCATION Section
        p = cell.add_paragraph()               # Create the paragraph
        run = p.add_run("EDUCATION")          # Add text
        run.bold = True                        # Bold
        run.font.size = Pt(12)                 # Font size
        run.font.color.rgb = grey(color_codes['1'])  # Dark grey color
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6) # Spacing after section header
        
        # self.add_section_header(cell, "EDUCATION")
        
        for edu in data['education']:
            p = cell.add_paragraph()
            run = p.add_run(edu['institution'])
            run.bold = True
            run.font.color.rgb = grey(color_codes['1'])
            run.font.size = Pt(11)
            
            run = p.add_run(f" --- ")
            run = p.add_run(edu['degree'])
            p.paragraph_format.space_after = Pt(0)
            run.font.color.rgb = grey(color_codes['1'])
            run.italic = True
            
            p = cell.add_paragraph()
            run = p.add_run(edu['year'])
            # run.bold = True
            run.font.color.rgb = grey(color_codes['1'])
            p.paragraph_format.space_before = Pt(0)
            run.font.size = Pt(10)
            
        # LANGUAGES Section
        p = cell.add_paragraph()               # Create the paragraph
        run = p.add_run("LANGUAGES")          # Add text
        run.bold = True                        # Bold
        run.font.size = Pt(12)                 # Font size
        run.font.color.rgb = grey(color_codes['1'])  # Dark grey color
        p.paragraph_format.space_after = Pt(6) # Spacing after section header
        
        # self.add_section_header(cell, "LANGUAGES")
        
        p = cell.add_paragraph()
        run = p.add_run(data['languages'])
        run.font.color.rgb = grey(color_codes['1'])
        p.paragraph_format.space_after = Pt(6)
    
    def add_section_header(self, cell, title):
        """Add a section header"""
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    def save(self, filename):
        """Save the document"""
        self.doc.save(filename)
        print(f"Resume saved as {filename}")
    
    # def save_as_pdf(self, docx_filename, pdf_filename=None):
    #     """Save the document as PDF"""
    #     if pdf_filename is None:
    #         pdf_filename = docx_filename.replace('.docx', '.pdf')
        
    #     # First save as docx
    #     self.save(docx_filename)
        
    #     # Convert to PDF
    #     convert(docx_filename, pdf_filename)
    #     print(f"PDF saved as {pdf_filename}")
    # def save_as_pdf(self, docx_filename, pdf_filename=None):
    #     """Save the document as PDF using LibreOffice"""
    #     if pdf_filename is None:
    #         pdf_filename = docx_filename.replace('.docx', '.pdf')
        
    #     # First save as docx
    #     self.save(docx_filename)
        
    #     # Convert to PDF using LibreOffice
    #     try:
    #         # Get the directory of the docx file
    #         output_dir = os.path.dirname(docx_filename) or '.'
            
    #         subprocess.run([
    #             'libreoffice', # or 'soffice',
    #             '--headless',
    #             '--convert-to', 'pdf',
    #             '--outdir', output_dir,
    #             docx_filename
    #         ],
    #         # capture_output=True, text=True, timeout=30, stderr=subprocess.DEVNULL)
    #         check=True)
            
    #         print(f"PDF saved as {pdf_filename}")
    #     except subprocess.CalledProcessError as e:
    #         print(f"Error converting to PDF: {e}")
    #         print("Make sure LibreOffice is installed: sudo apt install libreoffice")
    #     except FileNotFoundError:
    #         print("LibreOffice not found. Install it with: sudo apt install libreoffice")
    def save_as_pdf(self, docx_filename, pdf_filename=None):
        """
        Save the document as PDF preserving hyperlinks.
        Multiple methods in order of reliability.
        """
        if pdf_filename is None:
            pdf_filename = docx_filename.replace('.docx', '.pdf')
        
        # First save as docx
        self.save(docx_filename)
        
        # Method 1: Try unoconv (best for Linux, preserves links)
        try:
            subprocess.run([
                'unoconv',
                '-f', 'pdf',
                '-o', pdf_filename,
                docx_filename
            ], check=True, capture_output=True)
            print(f"PDF saved as {pdf_filename} (using unoconv)")
            return
        except (subprocess.CalledProcessError, FileNotFoundError):
            print("unoconv not available, trying next method...")


# Example usage
if __name__ == "__main__":
    # Sample data structure
    resume_data = {
        'name': 'Dafetite O. Ogaga',
        'title': 'Software and Automation Engineer',
        'location': 'Lagos, Nigeria',
        'phone': '+234 (0) 8038091572',
        'email': 'ogagadafetite@gmail.com',
        'portfolio_url': 'https://dafetiteogaga.github.io/dafetite/',
        'twitter_url': 'https://x.com/dafetite_ogaga',
        'linkedin_url': 'https://www.linkedin.com/in/ogagadafetite/',
        'github_url': 'https://github.com/DafetiteOgaga',
        'whatsapp_url': 'https://wa.link/67x975',
        
        'experience': [
            {
                'company': 'Freelance, Lagos',
                'role': 'Software and Automation Engineer',
                'dates': '2023 - PRESENT',
                'description': 'Provided end-to-end software and automation solutions to clients, leveraging technologies such as Python, JavaScript, and Django. Streamlined business workflows, automated manual processes, and built user-friendly applications across web and mobile platforms.'
            },
            {
                'company': 'Alx Africa, Online',
                'role': 'Software and Automation Engineer',
                'dates': '2021 - 2023',
                'description': 'Built a Betty linter wrapper compatible with Android. Created "Custom Commands," a modular Python CLI tool, and developed "Article Hive," a robust publishing platform featuring role-based access control, content management, and a user-friendly interface.'
            },
            {
                'company': 'Altaviz Support Limited, Lagos',
                'role': 'Software Engineer',
                'dates': '2018 - 2021',
                'description': 'Designed and implemented a workflow management system that streamlined internal communications, automated support processes, and improved task tracking across departments. Significantly enhanced operational efficiency and team collaboration.'
            }
        ],
        
        'projects': [
            {
                'name': 'FamousPropertiesNG',
                'type': 'Online marketplace',
                'description': 'An innovative decluttering and e-commerce platform that connects sellers looking to part with property items to buyers seeking quality deals. Designed to streamline the property resale process, the platform offers a seamless experience for discovering and purchasing items. Built with modern web technologies, FamousPropertiesNG emphasizes speed, scalability, and intuitive navigation, ensuring that users can browse listings, view detailed product information, and make orders with ease.'
            },
            {
                'name': 'StuddieBudie',
                'type': 'Online Education App and Utility',
                'description': 'An EdTech platform designed to enhance learning for students in basic and secondary schools. The application offers curriculum-aligned subject materials, auto-graded tests, and dynamic question pools. Teachers can contribute and scramble custom questions for download in up to 26 types, ideal for both digital learning and physical classroom use.'
            },
            {
                'name': 'Altaviz Support Limited',
                'type': 'Web and Mobile App (Role-based management workflow system)',
                'description': 'A robust role-based workflow management system built for Altaviz Support Limited. This web application streamlines operations across multiple departments, including Custodians, Engineers, Workshop staff, Help Desk, HR, and Supervisors. Developed using Django (backend), ReactJS (frontend), and MySQL (database), the platform is designed to be fast, scalable, and reliable, providing a seamless interface for day-to-day company operations.'
            },
            {
                'name': 'Article Hive',
                'type': 'An interactive webapp for lovers of articles',
                'description': 'An interactive web platform for readers, writers, and knowledge seekers. Article Hive functions as an online book club where users can register as members, publish articles, build an author profile, and grow their audience. Built with Django, Vanilla JavaScript, HTML, CSS, and MySQL, it offers a clean, intuitive user experience for both anonymous readers and contributing authors.'
            }
        ],
        
        'education': [
            {
                'institution': 'Federal University of Technology Akure, Ondo',
                'degree': 'B.Eng',
                'year': '2014'
            }
        ],
        
        'skills': [
            'Python, JavaScript(ES6), C, HTML5 and CSS',
            'Web Application, Server Database development',
            'Proficient with Linux and Windows CLIs',
            'UI/UX',
            'React.js, Django (Web Framework), and REST framework (DRF)',
            'Wireframes, Prototypes (Figma) designs',
            'Automation Scripting (Python, bash, batch, PowerShell)',
            'API, VSCode, Vim, Nano, and Emacs',
            'Git and GitHub'
        ],
        
        'awards': [
            {'title': 'Software Engineering (ALX)', 'year': '23/24'},
            {'title': 'Meta Front-End Developer (Specialization): Coursera (Meta)', 'year': '23/24'},
            {'title': 'Meta Back-End Developer (Specialization): Coursera (Meta)', 'year': '23/24'},
            {'title': 'IT Automation with Python (Specialization): Coursera (Google)', 'year': '23/24'},
            {'title': 'Google IT Support (Specialization): Coursera (Google)', 'year': '23/24'},
            {'title': 'Linux Tools and Libraries: Coursera (DARTMOUTH)', 'year': '23/24'}
        ],
        
        'languages': 'English, Yoruba'
    }
    
    # Icon paths (you'll need to provide these)
    icons = {
        'location': 'icons/location.png',
        'phone': 'icons/phone.png',
        'email': 'icons/email.png',
        'website': 'icons/website.png',
        'twitter': 'icons/twitter.png',
        'linkedin': 'icons/linkedin.png',
        'github': 'icons/github.png',
        'whatsapp': 'icons/whatsapp.png'
    }
    
    # Create resume
    generator = ResumeGenerator()
    generator.create_resume(resume_data, icons)
    
    # Save as both DOCX and PDF
    generator.save('Professional_Resume.docx')
    generator.save_as_pdf('Professional_Resume.docx', 'Professional_Resume.pdf')
    # generator.save_as_pdf('./Resume_new.docx', 'Resume_new.pdf')
